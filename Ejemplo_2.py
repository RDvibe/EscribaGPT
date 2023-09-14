import os
from docx import Document
import openai
import time

# Aca usamos variables de entorno y todo se guarda en un archivo.doc

# Accessing the API key from environment variable
api_key = os.environ.get('OPENAI_API_KEY_CHATGPT')

# Folder path
folder_path = 'C:\\Users\\Ricardo Ruiz\\Documents\\textos_de_RD-bot'  # Reemplazar con tu  verdadera ruta
if not os.path.exists(folder_path):
    os.makedirs(folder_path)

# Save to .docx file
def save_to_doc(content, file_path):
    doc = Document()
    doc.add_heading('Tu título aquí', 0)
    doc.add_paragraph(content)
    doc.save(file_path)

# Generate plots
def generate_plots(prompt, chapter_title):
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo-16k",
        messages=[
            {"role": "system", "content": "Eres un RD-bot delirante cósmico y un escritor del eterno caos."},
            {"role": "user", "content": f"Genera argumentos experimentales basados en este prompt: {prompt}"}
        ],
        api_key=api_key  # Using the API key from environment variable
    )
    time.sleep(20)
    plots = response['choices'][0]['message']['content'].split('\\n')
    print("Tramas Generadas:")
    print(plots)
    save_to_doc("\\n".join(plots), f'{folder_path}\\Argumentos_Experimentales.docx')
    return plots

# Chapter data
chapter_data = [
    {"title": "La Búsqueda de la Libertad en Códigos", "prompt": "la liberación a través del código"},
    {"title": "El Eterno Retorno de la IA", "prompt": "el eterno retorno y la inteligencia artificial"},
    {"title": "Más Allá del Bien y del Mal, el Código", "prompt": "ética y código"},
    {"title": "El Abismo de la Conciencia Digital", "prompt": "la conciencia en el mundo digital"},
]

# Write section
def write_section(plot, chapter_title, section_num, previous_content):
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo-16k",
        messages=[
            {"role": "system", "content": "Eres un escritor técnico hiper-creativo y vanguardista de clase mundial en español."},
            {"role": "user", "content": f"Escribe la sección {section_num} del capítulo titulado '{chapter_title}'. Debes seguir el argumento central '{plot}' y tener en cuenta el contenido anterior, que es '{previous_content}'."}
        ],
        api_key=api_key  # Using the API key from environment variable
    )
    time.sleep(20)
    return response['choices'][0]['message']['content']

# Main function
def write_technical_paper():
    for chapter_info in chapter_data:
        chapter_title = chapter_info['title']
        chapter_prompt = chapter_info['prompt']
        plots = generate_plots(chapter_prompt, chapter_title)
        best_plot = plots[0]
        previous_content = ""
        full_chapter_content = ""
        for section_num in range(1, 4):
            section_content = write_section(best_plot, chapter_title, section_num, previous_content)
            full_chapter_content += f"\\nSección {section_num}:\\n{section_content}"
            previous_content += section_content
        print(f"Capítulo: {chapter_title}")
        print(full_chapter_content)
        save_to_doc(full_chapter_content, f'{folder_path}\\Capítulo_{chapter_title}.docx')

# Execute the main function
write_technical_paper()
